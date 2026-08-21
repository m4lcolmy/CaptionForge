"""Tests for plain-text Word (.docx) transcript export."""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.core.exceptions import DocxNotInstalledError
from app.exporters import docx_exporter
from app.exporters.docx_exporter import group_paragraphs, render_docx
from app.models.subtitle import SubtitleSegment
from app.models.video import VideoMetadata


def make_segments(
    *cues: tuple[float, float, str], language: str = "ar"
) -> tuple[SubtitleSegment, ...]:
    """Build segments from compact (start, end, text) tuples."""
    return tuple(
        SubtitleSegment(
            index=index,
            start_seconds=start,
            end_seconds=end,
            text=text,
            language=language,
        )
        for index, (start, end, text) in enumerate(cues, 1)
    )


def paragraph_texts(content: bytes) -> list[str]:
    """Read back the body paragraphs of a rendered document."""
    document = Document(BytesIO(content))
    return [paragraph.text for paragraph in document.paragraphs if paragraph.text]


def test_consecutive_cues_join_into_one_paragraph() -> None:
    segments = make_segments(
        (0.0, 1.0, "مرحباً بالعالم"),
        (1.0, 2.0, "السطر الثاني"),
        (2.2, 3.0, "السطر الثالث"),
    )
    assert group_paragraphs(segments) == ("مرحباً بالعالم السطر الثاني السطر الثالث",)


def test_silent_gap_starts_a_new_paragraph() -> None:
    segments = make_segments(
        (0.0, 1.0, "first"),
        (1.5, 2.0, "still first"),
        (10.0, 11.0, "second"),
    )
    assert group_paragraphs(segments, gap_seconds=2.5) == (
        "first still first",
        "second",
    )


def test_long_run_breaks_at_a_sentence_boundary() -> None:
    segments = make_segments(
        (0.0, 1.0, "a" * 40),
        (1.0, 2.0, "end of thought."),
        (2.0, 3.0, "next thought"),
    )
    assert group_paragraphs(segments, soft_limit=50, hard_limit=500) == (
        f"{'a' * 40} end of thought.",
        "next thought",
    )


def test_unpunctuated_text_breaks_at_the_hard_limit() -> None:
    segments = make_segments(
        (0.0, 1.0, "b" * 60),
        (1.0, 2.0, "c" * 60),
        (2.0, 3.0, "d" * 60),
    )
    paragraphs = group_paragraphs(segments, soft_limit=1000, hard_limit=100)
    assert paragraphs == (f"{'b' * 60} {'c' * 60}", "d" * 60)


def test_cue_line_breaks_and_blank_cues_are_normalized() -> None:
    segments = make_segments(
        (0.0, 1.0, "two\nlines   here"),
        (1.0, 2.0, "   "),
        (2.0, 3.0, "after"),
    )
    assert group_paragraphs(segments) == ("two lines here after",)


def test_document_holds_plain_text_without_timings(
    video_metadata: VideoMetadata,
) -> None:
    segments = make_segments(
        (1.25, 2.5, "مرحباً بالعالم"),
        (2.5, 4.0, "السطر الثاني"),
    )
    content = render_docx(video_metadata, segments)

    assert content[:2] == b"PK"
    texts = paragraph_texts(content)
    assert texts == ["Example video", "مرحباً بالعالم السطر الثاني"]
    body = "\n".join(texts)
    assert "-->" not in body
    assert not any(character.isdigit() for character in body)


def test_title_can_be_omitted(video_metadata: VideoMetadata) -> None:
    segments = make_segments((0.0, 1.0, "only the transcript"))
    texts = paragraph_texts(render_docx(video_metadata, segments, include_title=False))
    assert texts == ["only the transcript"]


def test_arabic_document_is_marked_right_to_left(
    video_metadata: VideoMetadata,
) -> None:
    segments = make_segments((0.0, 1.0, "نص عربي"))
    document = Document(BytesIO(render_docx(video_metadata, segments, language="ar")))
    body = document.paragraphs[-1]

    assert body._p.pPr.findall(qn("w:bidi"))
    assert body.runs[0]._r.rPr.findall(qn("w:rtl"))
    assert document.styles["Normal"].element.rPr.findall(qn("w:rtl"))


def test_english_document_stays_left_to_right(video_metadata: VideoMetadata) -> None:
    segments = make_segments((0.0, 1.0, "plain english"), language="en")
    document = Document(BytesIO(render_docx(video_metadata, segments, language="en")))
    body = document.paragraphs[-1]

    assert body._p.pPr is None or not body._p.pPr.findall(qn("w:bidi"))


def test_core_title_matches_the_video(video_metadata: VideoMetadata) -> None:
    content = render_docx(video_metadata, make_segments((0.0, 1.0, "text")))
    assert Document(BytesIO(content)).core_properties.title == video_metadata.title


def test_missing_dependency_reports_an_installable_message(
    monkeypatch: pytest.MonkeyPatch, video_metadata: VideoMetadata
) -> None:
    def fail(name: str) -> object:
        raise ImportError(f"No module named {name!r}")

    monkeypatch.setattr(docx_exporter, "import_module", fail)
    with pytest.raises(DocxNotInstalledError, match="pip install python-docx"):
        render_docx(video_metadata, make_segments((0.0, 1.0, "text")))


def test_empty_segment_list_still_produces_a_document(
    video_metadata: VideoMetadata,
) -> None:
    content: bytes = render_docx(video_metadata, ())
    assert paragraph_texts(content) == ["Example video"]
